from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("GateFlow-v0.1-Instruction-Manual.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17324D"
MUTED = "5E6B78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CAUTION = "7A5A00"
RISK = "9B1C1C"
WHITE = "FFFFFF"


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color="C7D0D9"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def write_cell(cell, text, bold=False, color=INK, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_data_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_table_borders(table)
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        write_cell(cell, header, bold=True, color=DARK_BLUE, size=10)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            write_cell(cell, value)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, label, text, fill=LIGHT_GRAY, color=INK):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, "D5DDE5")
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, color=color, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5, color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True, color=INK)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest, color=INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, color=INK)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, color=INK)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, color=INK)
    return paragraph


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    set_run_font(run, size=8.5, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_para.paragraph_format.space_after = Pt(0)
    header_run = header_para.add_run("GateFlow v0.1 | Guard Instruction Manual")
    set_run_font(header_run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.paragraph_format.space_after = Pt(0)
    label_run = footer_para.add_run("Local prototype | Page ")
    set_run_font(label_run, size=8.5, color=MUTED)
    add_page_number(footer_para)


def add_cover(doc):
    for _ in range(8):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("OPERATIONAL REFERENCE")
    set_run_font(run, size=10, color=CAUTION, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("GateFlow v0.1")
    set_run_font(run, size=30, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run("Gate Scanner and Gate Activity Instruction Manual")
    set_run_font(run, size=15, color=DARK_BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(0)
    run = meta.add_run("Prototype version 0.1 | Guard, supervisor, and manager workflow")
    set_run_font(run, size=10.5, color=MUTED)

    doc.add_page_break()


def build_manual():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "Purpose and Scope")
    add_body(doc, "This manual explains how to use the GateFlow v0.1 prototype to record vehicle departures and returns at the gate. It covers the guard scanner workflow, supervisor approval, manager controls, search, and the audit record created by each action.")
    add_callout(doc, "Important", "GateFlow v0.1 is a local browser prototype. It saves its demonstration data in the current browser on the current device. It is not yet a shared production gate system or a replacement for site security policy.", fill=LIGHT_GRAY)

    add_heading(doc, "At a Glance", level=2)
    add_data_table(
        doc,
        ["Role", "Primary work in GateFlow"],
        [
            ("Gate guard", "Use the Scanner tab to record Vehicle OUT and Vehicle IN transactions."),
            ("Supervisor", "Approve an unauthorized active driver for the current day when an OUT is blocked."),
            ("Manager / admin", "Authorize or deauthorize drivers for today, review drivers, vehicles, locations, transactions, and audit events."),
        ],
        [2700, 6660],
    )

    add_heading(doc, "Daily Start")
    add_numbered(doc, "Open GateFlow and select the Scanner tab. The large Vehicle Out and Vehicle In buttons are the normal starting point for gate work.")
    add_numbered(doc, "Confirm that the date chip is correct and that the saved-status indicator reads Saved locally.")
    add_numbered(doc, "Review the compact counts for today: OUT, IN, and blocked attempts. The recent activity list shows the latest recorded transactions.")
    add_numbered(doc, "For this prototype, type an ID, VIN, plate, or barcode into the field. A future Zebra DataWedge keyboard-wedge scan will enter its payload into the focused field in the same way.")

    add_heading(doc, "Vehicle OUT")
    add_body(doc, "Use Vehicle OUT when a vehicle is leaving the facility. An active driver must either be authorized for today or be approved by an active supervisor before the transaction can be submitted.")

    add_heading(doc, "Record a Vehicle OUT", level=2)
    add_numbered(doc, "Tap Vehicle Out on the scanner home screen.")
    add_numbered(doc, "Scan or enter the Driver ID, then select Next. GateFlow checks that the driver exists, is active, and is authorized for the current day.")
    add_numbered(doc, "Scan or enter the vehicle VIN, plate, or barcode. GateFlow finds a vehicle by any of those values.")
    add_numbered(doc, "Choose the destination. A destination is required before the guard can continue.")
    add_numbered(doc, "Add a note when useful. The OUT flow also permits an optional photo preview for seals, load condition, visitors, or a guard note.")
    add_numbered(doc, "Review the driver, vehicle, destination, authorization method, note, and photo indicator. Select Submit OUT, then wait for the OUT submitted confirmation before releasing the vehicle.")

    add_callout(doc, "Photo note", "In v0.1, the guard can preview a selected photo and the transaction stores only that a photo was attached. The photo file itself is not persisted as a retrievable attachment yet.", fill="FFF8E8", color=CAUTION)

    add_heading(doc, "What Changes When a Vehicle Goes OUT", level=2)
    add_data_table(
        doc,
        ["System area", "Change after Submit OUT"],
        [
            ("Transaction record", "A new transaction with a unique GF ID, type OUT, timestamp/date, driver ID/name/company, VIN, barcode, plate, destination, note, supervisor ID when used, and photo-attached flag."),
            ("Vehicle record", "Vehicle status changes to out. Last location changes to the selected destination. Last transaction ID changes to the new OUT transaction."),
            ("Audit log", "A Vehicle OUT transaction event is added with the guard as actor and links to the transaction, driver, vehicle VIN, and any supervisor ID."),
            ("Scanner and dashboard", "Today OUT count rises. Recent activity, manager search results, vehicle status, and admin snapshot data reflect the new transaction."),
            ("Driver authorization", "No change for a normal already-authorized driver. A supervisor-approved OUT first authorizes the driver for the current day, then records the OUT."),
        ],
        [2850, 6510],
    )

    add_heading(doc, "OUT Blocked: Supervisor Approval", level=2)
    add_body(doc, "If the driver is active but not authorized today, GateFlow stops the OUT flow at the Supervisor required screen. The vehicle is not released and no OUT transaction is created until a valid supervisor approves.")
    add_numbered(doc, "Ask a valid supervisor to scan or enter their Supervisor ID.")
    add_numbered(doc, "Select Approve today. GateFlow returns to the vehicle scan step after a valid active supervisor is recognized.")
    add_numbered(doc, "Complete the remaining vehicle, destination, note/photo, review, and Submit OUT steps.")
    add_data_table(
        doc,
        ["Situation", "What GateFlow records"],
        [
            ("Driver lacks daily authorization", "Blocked OUT attempt audit event. No vehicle status or location change."),
            ("Valid supervisor approves", "Driver authorized for today, Supervisor override audit event, and Driver authorized audit event. The actual OUT transaction is still recorded separately after submission."),
            ("Invalid supervisor credential", "Invalid supervisor credential audit event. The OUT remains blocked."),
            ("Unknown or inactive driver", "Blocked OUT attempt audit event. The OUT remains blocked."),
        ],
        [3050, 6310],
    )

    add_heading(doc, "Vehicle IN")
    add_body(doc, "Use Vehicle IN when a vehicle returns to the facility. The driver must exist and be active. Daily OUT authorization is not required for a return in this prototype, but inactive drivers are blocked.")

    add_heading(doc, "Record a Vehicle IN", level=2)
    add_numbered(doc, "Tap Vehicle In on the scanner home screen.")
    add_numbered(doc, "Scan or enter the returning Driver ID. The scanner confirms the driver is active.")
    add_numbered(doc, "Scan or enter the vehicle VIN, plate, or barcode.")
    add_numbered(doc, "Choose the location the vehicle is returning from. This is required.")
    add_numbered(doc, "Add an optional condition, mileage, damage, or guard note. Review the transaction and select Submit IN. Wait for the IN submitted confirmation.")

    add_heading(doc, "What Changes When a Vehicle Comes IN", level=2)
    add_data_table(
        doc,
        ["System area", "Change after Submit IN"],
        [
            ("Transaction record", "A new transaction with a unique GF ID, type IN, timestamp/date, driver ID/name/company, VIN, barcode, plate, from location, note, no supervisor ID, and no persisted photo."),
            ("Vehicle record", "Vehicle status changes to in. Last location is set to North Gate 4 in the current v0.1 prototype. Last transaction ID changes to the new IN transaction."),
            ("Audit log", "A Vehicle IN transaction event is added with the guard as actor and links to the transaction, driver, and vehicle VIN."),
            ("Scanner and dashboard", "Today IN count rises. Recent activity, manager search results, vehicle status, and admin snapshot data reflect the return."),
            ("Driver authorization", "No authorization setting changes are made by a normal IN transaction."),
        ],
        [2850, 6510],
    )

    add_heading(doc, "Exceptions and Corrections")
    add_data_table(
        doc,
        ["Screen result", "Guard action"],
        [
            ("Driver not found", "Re-scan or correct the driver ID. In the OUT flow, the attempt is added to the audit log as blocked. No transaction is created."),
            ("Vehicle not found", "Re-scan the VIN, plate, or barcode. No transaction is created and vehicle data is unchanged."),
            ("Driver inactive", "Do not process the transaction. OUT and IN attempts for an inactive known driver are blocked and added to the audit log."),
            ("No daily OUT authorization", "Use the supervisor approval flow. Do not release the vehicle on the basis of the blocked screen alone."),
            ("Need to stop a flow", "Use Cancel OUT or Cancel IN, or return to the scanner home screen. The unfinished form data is cleared and no transaction is saved."),
        ],
        [2800, 6560],
    )

    add_heading(doc, "Admin and Manager Tasks")
    add_heading(doc, "Authorize Drivers for Today", level=2)
    add_body(doc, "Open the Admin tab. In the Drivers list, use the authorize or deauthorize control for the selected driver.")
    add_bullet(doc, "Authorize sets the driver authorization date to today and adds a Driver authorized audit event with Manager as the actor.")
    add_bullet(doc, "Deauthorize clears the authorization date and adds a Driver deauthorized audit event with Manager as the actor.")
    add_bullet(doc, "The Admin tab also lists vehicles and approved locations, and its snapshot shows transactions, overrides, vehicles out, and active drivers.")

    add_heading(doc, "Find Gate Activity", level=2)
    add_body(doc, "Open the Search tab to find recorded transactions. Filters accept VIN or partial VIN, plate or partial plate, driver text, date, location, and IN or OUT. The results table shows time, direction, vehicle, VIN, driver, location, and note.")

    add_heading(doc, "Review the Audit Log", level=2)
    add_body(doc, "Open the Audit tab to review security events. Filter by IN/OUT transaction, blocked attempt, supervisor override, or authorization change. Text search can match a supervisor, driver, VIN, or note.")

    add_heading(doc, "End of Shift")
    add_numbered(doc, "Review the Admin snapshot and identify vehicles still marked out.")
    add_numbered(doc, "Use Search to review transactions for the current date, location, or direction when reconciling activity.")
    add_numbered(doc, "Use the Audit log to confirm blocked attempts, supervisor overrides, and authorization changes that need follow-up.")
    add_numbered(doc, "Do not use Reset demo during an active demonstration unless you intend to discard all locally saved prototype data and restore the supplied seed data.")

    add_heading(doc, "Prototype Limits and Production Path")
    add_data_table(
        doc,
        ["Current v0.1 behavior", "Production implication"],
        [
            ("Saved locally in browser storage", "A production rollout needs Supabase/Postgres or customer-hosted enterprise storage with shared, durable records."),
            ("Typed fields and simulated scans", "Zebra DataWedge keyboard wedge or Android Intent input can route real scanner payloads into these focused fields."),
            ("No live sign-in or role enforcement", "Production requires role-based permissions for guards, supervisors, managers, and auditors."),
            ("Photo indicator only", "Production requires camera capture, secure file storage, retention policy, and retrieval controls."),
            ("No sync engine", "An offline-first mobile build needs queued writes, conflict handling, and sync status."),
            ("Browser prototype", "A later Zebra Enterprise Browser or native Android application can use the same transaction rules with device APIs."),
        ],
        [3050, 6310],
    )

    add_heading(doc, "Demo IDs Included in the Prototype")
    add_data_table(
        doc,
        ["Value", "Use in the scanner"],
        [
            ("D-1027", "Authorized driver sample for a normal OUT flow."),
            ("D-2033", "Active driver sample that requires supervisor approval before OUT."),
            ("SUP-1001", "Valid supervisor sample for the approval screen."),
            ("YARD-104", "Vehicle plate/barcode sample."),
        ],
        [2500, 6860],
    )

    add_callout(doc, "Release rule", "For an OUT, treat the confirmation screen as the record of completion. A scan or supervisor approval alone does not create the vehicle OUT transaction.", fill="FFF0F0", color=RISK)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_manual()
